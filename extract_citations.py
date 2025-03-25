import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import docx
from docx import Document
from docx.shared import Pt
from eyecite import get_citations
from eyecite.models import FullCaseCitation, ShortCaseCitation
import pyperclip  # make sure to install this with `pip install pyperclip`
import os


def extract_citations_from_docx(filepath):
    """Reads a .docx file and returns citations using eyecite."""
    doc = docx.Document(filepath)
    text = " ".join(para.text for para in doc.paragraphs)
    return get_citations(text)


def format_citation(citation):
    """Formats a citation into a human-readable legal format string."""
    try:
        matched = citation.matched_text()

        if isinstance(citation, FullCaseCitation):
            plaintiff = citation.metadata.plaintiff or ""
            defendant = citation.metadata.defendant or ""
            case_name = f"{plaintiff} v. {defendant}" if plaintiff and defendant else matched
            pin_cite = f", {citation.metadata.pin_cite}" if citation.metadata.pin_cite else ""
            year = f" ({citation.metadata.year})" if citation.metadata.year else ""
            return f"{case_name}, {matched}{pin_cite}{year}"

        elif isinstance(citation, ShortCaseCitation):
            pin_cite = f" at {citation.metadata.pin_cite}" if citation.metadata.pin_cite else ""
            return f"{matched}{pin_cite}"

        else:
            return str(matched)

    except Exception as e:
        return f"[Formatting error: {str(e)}]"


def load_file():
    """Handles file selection and displays formatted citations in the UI."""
    global original_doc_path, formatted_citations
    filepath = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if not filepath:
        return

    try:
        citations = extract_citations_from_docx(filepath)
        formatted = [format_citation(c) for c in citations if isinstance(c, (FullCaseCitation, ShortCaseCitation))]
        formatted_citations = sorted(set(str(f) for f in formatted))

        original_doc_path = filepath

        text_area.config(state=tk.NORMAL)
        text_area.delete('1.0', tk.END)
        text_area.insert(tk.END, "\n".join(formatted_citations))
        text_area.config(state=tk.DISABLED)

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred:\n{str(e)}")


def export_to_txt():
    """Exports citations to a .txt file."""
    if not formatted_citations:
        messagebox.showinfo("No data", "No citations to export.")
        return

    output_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text file", "*.txt")])
    if not output_path:
        return

    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(formatted_citations))
        messagebox.showinfo("Success", f"Exported to {output_path}")
    except Exception as e:
        messagebox.showerror("Export Failed", str(e))


def export_to_docx():
    """Exports citations to a new .docx file."""
    if not formatted_citations:
        messagebox.showinfo("No data", "No citations to export.")
        return

    output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if not output_path:
        return

    try:
        doc = Document()
        doc.add_heading("Table of Authorities", 0)
        for citation in formatted_citations:
            para = doc.add_paragraph(citation)
            para.style.font.size = Pt(11)
        doc.save(output_path)
        messagebox.showinfo("Success", f"Exported to {output_path}")
    except Exception as e:
        messagebox.showerror("Export Failed", str(e))


def copy_to_clipboard():
    """Copies the citations to clipboard."""
    if not formatted_citations:
        messagebox.showinfo("No data", "No citations to copy.")
        return
    pyperclip.copy("\n".join(formatted_citations))
    messagebox.showinfo("Copied", "Citations copied to clipboard.")


def insert_into_original_doc():
    """Appends citations to the end of the original .docx file as TOA."""
    if not formatted_citations:
        messagebox.showinfo("No data", "No citations to insert.")
        return
    if not original_doc_path:
        messagebox.showinfo("No file", "Original document not found.")
        return

    try:
        doc = Document(original_doc_path)
        doc.add_page_break()
        doc.add_heading("Table of Authorities", level=1)
        for citation in formatted_citations:
            para = doc.add_paragraph(citation)
            para.style.font.size = Pt(11)

        output_path = os.path.splitext(original_doc_path)[0] + "_with_TOA.docx"
        doc.save(output_path)
        messagebox.showinfo("Inserted", f"Citations added to: {output_path}")
    except Exception as e:
        messagebox.showerror("Insert Failed", str(e))


def create_app():
    """Creates and runs the Tkinter app."""
    app = tk.Tk()
    app.title("Legal Citation Extractor")
    app.geometry("900x600")

    global text_area
    text_area = scrolledtext.ScrolledText(app, width=100, height=25, wrap=tk.WORD)
    text_area.pack(padx=20, pady=20)
    text_area.config(state=tk.DISABLED)

    button_frame = tk.Frame(app)
    button_frame.pack(pady=10)

    tk.Button(button_frame, text="Load .docx", command=load_file).grid(row=0, column=0, padx=5)
    tk.Button(button_frame, text="Export to .txt", command=export_to_txt).grid(row=0, column=1, padx=5)
    tk.Button(button_frame, text="Export to .docx", command=export_to_docx).grid(row=0, column=2, padx=5)
    tk.Button(button_frame, text="Copy to Clipboard", command=copy_to_clipboard).grid(row=0, column=3, padx=5)
    tk.Button(button_frame, text="Insert TOA into Original Doc", command=insert_into_original_doc).grid(row=0, column=4, padx=5)

    return app


# Globals to hold state
formatted_citations = []
original_doc_path = None


if __name__ == "__main__":
    app = create_app()
    app.mainloop()
